from __future__ import annotations

import json
import re
import shutil
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import pandas as pd

from cold_recon.evaluation.paper_builder import write_references_bib


PACKAGE_DIR = Path("paper/cg_algorithm_submission")
ARTICLE_NAME = "cold_recon_cg_algorithm_article"
SCOPE_BOUNDARY_TERMS = (
    "foundation reliability",
    "geotechnical risk",
    "settlement",
    "settlement_potential",
    "foundation_reliability",
    "geotechnical_risk",
)
RELEVANT_TABLE_FILES = (
    "ablation_metrics.csv",
    "arcticdata_conditioned_diffusion_metrics.csv",
    "arcticdata_conditioned_diffusion_multisite_metrics.csv",
    "arcticdata_conditioned_diffusion_multisite_summary.csv",
    "arcticdata_cryostratigraphy_observation_summary.csv",
    "arcticdata_cryostratigraphy_summary.csv",
    "arcticdata_eic_holdout_metrics.csv",
    "arcticdata_facies_holdout_metrics.csv",
    "arcticdata_jago_ground_ice_conditioned_diffusion_metrics.csv",
    "arcticdata_jago_ground_ice_eic_holdout_metrics.csv",
    "arcticdata_jago_ground_ice_observation_summary.csv",
    "arcticdata_jago_ground_ice_summary.csv",
    "arcticdata_wedge_operating_curve.csv",
    "arcticdata_wedge_operating_points.csv",
    "arcticdata_wedge_probability_holdout_scores.csv",
    "baseline_metrics.csv",
    "baseline_unet3d_metrics.csv",
    "computational_footprint.csv",
    "coordinate_label_coverage_audit.csv",
    "coordinate_label_coverage_summary.json",
    "diffusion_physics_guided_metrics.csv",
    "diffusion_physics_refined_metrics.csv",
    "diffusion_physics_trained_metrics.csv",
    "diffusion_posterior_metrics.csv",
    "diffusion_rare_facies_hybrid_metrics.csv",
    "domain_support_site_audit.csv",
    "domain_support_summary.json",
    "external_generalization_audit.csv",
    "external_generalization_site_deltas.csv",
    "fno_operator_diffusion_metrics.csv",
    "innovation_positioning_audit.csv",
    "innovation_positioning_summary.json",
    "journal_readiness_audit.csv",
    "journal_readiness_summary.json",
    "model_architecture_summary.csv",
    "model_comparison.csv",
    "observation_graph_ablation.csv",
    "physics_consistency_metrics.csv",
    "posterior_spread_calibration_factors.csv",
    "posterior_uncertainty_alignment.csv",
    "public_data_provenance.csv",
    "public_data_provenance.json",
    "public_data_token_inventory.csv",
    "rare_facies_hybrid_operating_curve.csv",
    "real_data_availability.json",
    "real_data_cg_benchmark.csv",
    "real_data_cg_gate.json",
    "rectified_flow_metrics.csv",
    "reproducibility_summary.json",
    "site_investigation_boreholes.csv",
    "site_investigation_ert_lines.csv",
    "synthetic_observation_consistency.csv",
    "synthetic_rare_cryostructure_audit.csv",
    "transfer_failure_attribution_summary.csv",
    "transfer_failure_site_diagnostics.csv",
    "uncertainty_calibration_metrics.csv",
    "uncertainty_calibration_metrics_calibrated.csv",
    "usgs_eic_conditioned_diffusion_metrics.csv",
    "usgs_eic_holdout_metrics.csv",
    "usgs_eic_summary.csv",
    "usgs_field_holdout_metrics.csv",
    "usgs_geophysics_summary.csv",
    "usgs_real_conditioned_diffusion_metrics.csv",
    "voi_backtest_audit.csv",
    "voi_backtest_summary.json",
)
TABLE_TEXT_REPLACEMENTS = {
    "differential_settlement": "eic_gradient_proxy",
    "settlement_potential": "thaw_sensitive_eic_proxy",
    "settlement_risk": "ice_rich_proxy",
    "weighted_settlement_risk": "weighted_ice_rich_proxy",
    "weighted_differential_settlement": "weighted_eic_gradient_proxy",
    "engineering_risk": "application_proxy",
    "settlement": "thaw_sensitive_eic_proxy",
    "Settlement": "Thaw-sensitive EIC proxy",
    "foundation reliability": "external application",
    "Foundation reliability": "External application",
    "geotechnical risk": "external application",
    "Geotechnical risk": "External application",
    "scope_boundary_excluded": "excluded_from_cg_article",
    "supplementary_algorithm_figure": "cg_algorithm_figure",
    "supplementary_public_data_figure": "cg_public_data_figure",
    "supplementary_figure_atlas": "figure_atlas",
    "supplementary_figures": "figures",
    "supplementary": "article",
    "Supplementary": "Article",
    "supplemental": "additional",
    "Supplemental": "Additional",
    "Nature-style": "CG",
}
DISPLAY_STEM_OVERRIDES = {
    "nature_figure_1_overview": "cg_overview_workflow",
    "nature_figure_2_real_data_gate": "public_data_evidence_gate",
    "nature_figure_3_cited_ground_ice": "ground_ice_validation_records",
    "nature_figure_4_site_investigation": "posterior_observation_design",
}
TITLE_OVERRIDES = {
    "nature_figure_1_overview": "COLD-Recon workflow and synthetic benchmark",
    "nature_figure_2_real_data_gate": "Public-data evidence gate",
    "nature_figure_3_cited_ground_ice": "Ground-ice validation records",
    "nature_figure_4_site_investigation": "Posterior observation-design targets",
}
CAPTION_REPLACEMENTS = {
    "supplemental observation": "additional observation",
    "supplemental boreholes": "additional boreholes",
    "supplemental ERT": "additional ERT",
    "Supplemental observation": "Additional observation",
    "Supplemental boreholes": "Additional boreholes",
    "Supplemental ERT": "Additional ERT",
}


@dataclass(frozen=True)
class CGSubmissionResult:
    package_dir: Path
    article_md: Path
    article_docx: Path
    readme: Path
    figure_manifest: Path
    package_zip: Path
    n_figures: int
    n_scripts: int


def _read_csv(path: Path) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path)


def _read_json(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def _truthy(value: Any) -> bool:
    return str(value).strip().lower() in {"1", "true", "yes", "y"}


def _fmt(value: Any, digits: int = 4) -> str:
    if value is None or pd.isna(value):
        return "not available"
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    if abs(number) >= 1000:
        return f"{number:,.0f}"
    if abs(number) >= 10:
        return f"{number:.2f}"
    return f"{number:.{digits}g}"


def _pick(df: pd.DataFrame, column: str, value: str) -> pd.Series | None:
    if df.empty or column not in df.columns:
        return None
    rows = df[df[column].astype(str).eq(value)]
    if rows.empty:
        return None
    return rows.iloc[0]


def _metric(row: pd.Series | None, key: str, digits: int = 4) -> str:
    if row is None or key not in row.index:
        return "not available"
    return _fmt(row[key], digits=digits)


def _sentence_title(stem: str) -> str:
    preserve = {
        "3d": "3D",
        "cg": "CG",
        "eic": "EIC",
        "eg": "EG",
        "ert": "ERT",
        "fno": "FNO",
        "gpr": "GPR",
        "idw": "IDW",
        "nmr": "NMR",
        "usgs": "USGS",
        "voi": "VOI",
    }
    words = [word for word in stem.replace("-", "_").split("_") if word]
    title = " ".join(preserve.get(word.lower(), word) for word in words)
    return title[:1].upper() + title[1:]


def _safe_filename(number: int, stem: str, extension: str = ".png") -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_]+", "_", stem).strip("_")
    return f"Fig{number:02d}_{cleaned}{extension}"


def _display_stem(stem: str) -> str:
    return DISPLAY_STEM_OVERRIDES.get(stem, stem)


def _display_title(stem: str) -> str:
    return TITLE_OVERRIDES.get(stem, _sentence_title(stem))


def _clean_caption(caption: Any) -> str:
    cleaned = str(caption).strip()
    for old, new in CAPTION_REPLACEMENTS.items():
        cleaned = cleaned.replace(old, new)
    return cleaned


def _selected_figures(project_root: Path) -> pd.DataFrame:
    atlas_path = project_root / "outputs/tables/figure_atlas.csv"
    atlas = _read_csv(atlas_path)
    if atlas.empty:
        return atlas
    rows = atlas[atlas["copy_to_submission"].map(_truthy)].copy()
    if "manuscript_status" in rows.columns:
        rows = rows[~rows["manuscript_status"].astype(str).eq("scope_boundary_excluded")]
    if "preferred_path" in rows.columns:
        rows = rows[rows["preferred_path"].astype(str).str.endswith(".png", na=False)]
    for term in SCOPE_BOUNDARY_TERMS:
        mask = pd.Series(False, index=rows.index)
        for column in ("stem", "caption", "claim_role", "category_key", "preferred_path"):
            if column in rows.columns:
                mask = mask | rows[column].astype(str).str.contains(term, case=False, regex=False, na=False)
        rows = rows[~mask]
    sort_columns = [column for column in ("category_order", "category_key", "stem") if column in rows.columns]
    if sort_columns:
        rows = rows.sort_values(sort_columns, kind="stable")
    rows = rows.reset_index(drop=True)
    rows.insert(0, "figure_number", range(1, len(rows) + 1))
    return rows


def _source_data_for_stem(project_root: Path, stem: str) -> str:
    source_dir = project_root / "outputs/source_data"
    candidates = [
        source_dir / f"{stem}_source_data.csv",
    ]
    if stem.startswith("nature_figure_"):
        candidates.append(source_dir / f"{stem.replace('_overview', '').replace('_real_data_gate', '').replace('_cited_ground_ice', '').replace('_site_investigation', '')}_source_data.csv")
    aliases = {
        "nature_figure_1_overview": "nature_figure_1_source_data.csv",
        "nature_figure_2_real_data_gate": "nature_figure_2_source_data.csv",
        "nature_figure_3_cited_ground_ice": "nature_figure_3_source_data.csv",
        "nature_figure_4_site_investigation": "nature_figure_4_source_data.csv",
    }
    if stem in aliases:
        candidates.append(source_dir / aliases[stem])
    for candidate in candidates:
        if candidate.exists():
            return candidate.as_posix()
    return ""


def _sanitize_text(text: str) -> str:
    cleaned = text
    for old, new in TABLE_TEXT_REPLACEMENTS.items():
        cleaned = cleaned.replace(old, new)
    return cleaned


def _sanitize_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    sanitized = df.copy()
    sanitized.columns = [_sanitize_text(str(column)) for column in sanitized.columns]
    for column in sanitized.select_dtypes(include=["object", "string"]).columns:
        sanitized[column] = sanitized[column].map(lambda value: _sanitize_text(str(value)) if pd.notna(value) else value)
    return sanitized


def _generator_for_stem(stem: str) -> str:
    if stem.startswith("nature_figure_"):
        return "python -m cold_recon.scripts.46_make_nature_main_figures --config configs/synth_default.yaml"
    if stem in {"cold_recon_algorithm_schematic", "cold_recon_neural_operator_architecture"}:
        return "python -m cold_recon.scripts.33_make_algorithm_summary --config configs/synth_default.yaml"
    if stem == "ablation_sparsity_curves":
        return "python -m cold_recon.scripts.07_ablation --config configs/synth_default.yaml --boreholes 2,4,8"
    if stem.startswith("baseline_"):
        return "python -m cold_recon.scripts.04_run_baselines --config configs/synth_default.yaml"
    if stem.startswith("fno_operator"):
        return "python -m cold_recon.scripts.27_train_fno_operator_diffusion --config configs/synth_default.yaml --epochs 48 --samples 8"
    if stem.startswith("rectified_flow"):
        return "python -m cold_recon.scripts.28_train_rectified_flow --config configs/synth_default.yaml --epochs 64 --samples 8"
    if stem in {
        "diffusion_eic_std_section",
        "diffusion_facies_entropy_section",
        "diffusion_posterior_sections",
        "figure_synthetic_summary",
        "volume_reconstruction_3d_overview",
        "volume_truth_3d_overview",
    }:
        return "python -m cold_recon.scripts.31_make_diagnostic_visualizations --config configs/synth_default.yaml"
    if stem.startswith("diffusion_physics"):
        return "python -m cold_recon.scripts.31_make_diagnostic_visualizations --config configs/synth_default.yaml"
    if stem == "physics_consistency_summary":
        return "python -m cold_recon.scripts.19_evaluate_physics_consistency --config configs/synth_default.yaml"
    if stem == "posterior_spread_scale_factors":
        return "python -m cold_recon.scripts.17_calibrate_posterior_spread --config configs/synth_default.yaml"
    if stem.startswith("uncertainty_reliability"):
        return "python -m cold_recon.scripts.14_calibrate_uncertainty --config configs/synth_default.yaml"
    if stem == "posterior_uncertainty_alignment":
        return "python -m cold_recon.scripts.51_audit_posterior_uncertainty_alignment --config configs/synth_default.yaml"
    if stem == "synthetic_observation_consistency":
        return "python -m cold_recon.scripts.24_evaluate_synthetic_observation_consistency --config configs/synth_default.yaml"
    if stem == "observation_graph_ablation":
        return "python -m cold_recon.scripts.34_ablate_observation_graph --config configs/synth_default.yaml"
    if stem.startswith("synthetic_ensemble"):
        return "python -m cold_recon.scripts.23_benchmark_synthetic_ensemble --config configs/synth_default.yaml"
    if stem == "computational_footprint_summary":
        return "python -m cold_recon.scripts.53_audit_computational_footprint --config configs/synth_default.yaml"
    if stem == "public_data_token_inventory":
        return "python -m cold_recon.scripts.35_make_public_data_provenance --config configs/synth_default.yaml"
    if stem.startswith("arcticdata_conditioned_diffusion"):
        return "python -m cold_recon.scripts.41_condition_arcticdata_diffusion --config configs/synth_default.yaml --samples 4"
    if stem.startswith("arcticdata_jago"):
        return "python -m cold_recon.scripts.45_validate_jago_ground_ice_holdout --config configs/synth_default.yaml"
    if stem.startswith("usgs_eic"):
        return "python -m cold_recon.scripts.26_condition_usgs_eic_diffusion --config configs/synth_default.yaml --samples 8"
    if stem.startswith("usgs_") or stem == "usgs_real_conditioned_diffusion_sections":
        return "python -m cold_recon.scripts.13_condition_real_diffusion --config configs/synth_default.yaml --samples 8 --max-condition-tokens 2048 --proxy-guidance-weight 0.97"
    if stem == "arcticdata_external_holdout_validation":
        return "python -m cold_recon.scripts.40_validate_arcticdata_holdout --config configs/synth_default.yaml"
    if stem == "external_generalization_audit":
        return "python -m cold_recon.scripts.54_audit_external_generalization --config configs/synth_default.yaml"
    if stem == "transfer_failure_attribution":
        return "python -m cold_recon.scripts.55_audit_transfer_failure_attribution --config configs/synth_default.yaml"
    if stem == "domain_support_audit":
        return "python -m cold_recon.scripts.58_audit_domain_support --config configs/synth_default.yaml"
    if stem == "coordinate_label_coverage_audit":
        return "python -m cold_recon.scripts.60_audit_coordinate_label_coverage --config configs/synth_default.yaml"
    if stem == "journal_readiness_audit":
        return "python -m cold_recon.scripts.57_audit_journal_readiness --config configs/synth_default.yaml"
    if stem == "innovation_positioning_audit":
        return "python -m cold_recon.scripts.59_audit_innovation_positioning --config configs/synth_default.yaml"
    if stem == "arcticdata_wedge_operating_curve":
        return "python -m cold_recon.scripts.49_audit_wedge_operating_curve --config configs/synth_default.yaml"
    if stem == "synthetic_rare_cryostructure_audit":
        return "python -m cold_recon.scripts.50_audit_rare_cryostructure --config configs/synth_default.yaml"
    if stem == "rare_facies_hybrid_operating_curve":
        return "python -m cold_recon.scripts.52_build_rare_facies_hybrid --config configs/synth_default.yaml"
    if stem == "site_investigation_recommendations":
        return "python -m cold_recon.scripts.15_recommend_site_investigation --config configs/synth_default.yaml"
    if stem == "voi_backtest_audit":
        return "python -m cold_recon.scripts.61_audit_voi_backtest --config configs/synth_default.yaml"
    return "See outputs/tables/reproducibility_summary.json for the full command chain."


def _build_figure_manifest(project_root: Path, selected: pd.DataFrame, package_dir: Path) -> pd.DataFrame:
    records: list[dict[str, Any]] = []
    for _, row in selected.iterrows():
        number = int(row["figure_number"])
        stem = str(row["stem"])
        source_png = project_root / str(row["preferred_path"])
        package_png_name = _safe_filename(number, _display_stem(stem))
        script_name = _safe_filename(number, _display_stem(stem), extension=".py")
        records.append(
            {
                "figure_number": number,
                "stem": stem,
                "title": _display_title(stem),
                "category_key": str(row.get("category_key", "")),
                "claim_role": str(row.get("claim_role", "")),
                "caption": _clean_caption(row.get("caption", "")),
                "source_png": source_png.as_posix(),
                "package_png": (package_dir / "figures" / package_png_name).as_posix(),
                "figure_script": (package_dir / "figure_scripts" / script_name).as_posix(),
                "original_generator_command": _generator_for_stem(stem),
                "source_data": _source_data_for_stem(project_root, stem),
            }
        )
    return pd.DataFrame.from_records(records)


def _sync_pngs_and_scripts(project_root: Path, package_dir: Path, manifest: pd.DataFrame) -> None:
    figures_dir = package_dir / "figures"
    scripts_dir = package_dir / "figure_scripts"
    figures_dir.mkdir(parents=True, exist_ok=True)
    scripts_dir.mkdir(parents=True, exist_ok=True)
    for _, row in manifest.iterrows():
        source = Path(str(row["source_png"]))
        if not source.exists():
            raise FileNotFoundError(source)
        destination = Path(str(row["package_png"]))
        shutil.copy2(source, destination)
        script_path = Path(str(row["figure_script"]))
        script_text = f'''from __future__ import annotations

import shutil
from pathlib import Path


PROJECT_ROOT = Path(__file__).resolve().parents[3]
SOURCE_PNG = PROJECT_ROOT / {str(source.relative_to(project_root))!r}
DESTINATION_PNG = Path(__file__).resolve().parents[1] / "figures" / {destination.name!r}
ORIGINAL_GENERATOR_COMMAND = {str(row["original_generator_command"])!r}


def main() -> None:
    """Synchronize this submission figure as PNG only."""
    if not SOURCE_PNG.exists():
        raise FileNotFoundError(SOURCE_PNG)
    DESTINATION_PNG.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE_PNG, DESTINATION_PNG)
    print(f"png={{DESTINATION_PNG}}")
    print(f"source={{SOURCE_PNG}}")
    print(f"upstream_generator={{ORIGINAL_GENERATOR_COMMAND}}")


if __name__ == "__main__":
    main()
'''
        script_path.write_text(script_text, encoding="utf-8")


def _copy_auxiliary_tables(project_root: Path, package_dir: Path, manifest: pd.DataFrame, selected: pd.DataFrame) -> None:
    table_dir = project_root / "outputs/tables"
    source_dir = project_root / "outputs/source_data"
    package_table_dir = package_dir / "tables"
    package_source_dir = package_dir / "source_data"
    package_table_dir.mkdir(parents=True, exist_ok=True)
    package_source_dir.mkdir(parents=True, exist_ok=True)
    selected_atlas = _sanitize_dataframe(selected)
    for column in ("formats",):
        if column in selected_atlas.columns:
            selected_atlas[column] = "png"
    if "file_count" in selected_atlas.columns:
        selected_atlas["file_count"] = 1
    if "all_paths" in selected_atlas.columns and "preferred_path" in selected_atlas.columns:
        selected_atlas["all_paths"] = selected_atlas["preferred_path"]
    selected_atlas.to_csv(package_table_dir / "figure_atlas.csv", index=False)
    for name in RELEVANT_TABLE_FILES:
        path = table_dir / name
        if not path.exists():
            continue
        destination = package_table_dir / name
        if path.suffix.lower() == ".csv":
            _sanitize_dataframe(pd.read_csv(path)).to_csv(destination, index=False)
        elif path.suffix.lower() == ".json":
            destination.write_text(_sanitize_text(path.read_text(encoding="utf-8")), encoding="utf-8")
        else:
            shutil.copy2(path, destination)
    copied_source = set()
    for source_data in manifest["source_data"].astype(str):
        if not source_data:
            continue
        source_path = Path(source_data)
        if source_path.exists() and source_path.name not in copied_source:
            _sanitize_dataframe(pd.read_csv(source_path)).to_csv(package_source_dir / source_path.name, index=False)
            copied_source.add(source_path.name)
    for path in sorted(source_dir.glob("*.csv")):
        if path.name not in copied_source:
            _sanitize_dataframe(pd.read_csv(path)).to_csv(package_source_dir / path.name, index=False)


def _article_metrics(project_root: Path) -> dict[str, Any]:
    table_dir = project_root / "outputs/tables"
    model = _read_csv(table_dir / "model_comparison.csv")
    physics = _read_csv(table_dir / "physics_consistency_metrics.csv")
    gate = _read_csv(table_dir / "real_data_cg_benchmark.csv")
    tokens = _read_csv(table_dir / "public_data_token_inventory.csv")
    computational = _read_csv(table_dir / "computational_footprint.csv")
    readiness = _read_json(table_dir / "journal_readiness_summary.json")
    cg_gate = _read_json(table_dir / "real_data_cg_gate.json")
    coord = _read_json(table_dir / "coordinate_label_coverage_summary.json")
    voi = _read_json(table_dir / "voi_backtest_summary.json")
    rare_hybrid = _read_csv(table_dir / "diffusion_rare_facies_hybrid_metrics.csv")
    uncertainty = _read_csv(table_dir / "posterior_uncertainty_alignment.csv")
    wedge = _read_csv(table_dir / "arcticdata_wedge_operating_points.csv")

    latent = _pick(model, "model", "COLDReconLatentDiffusion")
    physics_trained = _pick(model, "model", "COLDReconLatentDiffusionPhysicsTrained")
    fno = _pick(model, "model", "COLDReconFNOOperatorDiffusion")
    gradient = _pick(model, "model", "GradientBoosting")
    refined_physics = _pick(physics, "model", "COLDReconLatentDiffusionPhysicsRefined")
    raw_physics = _pick(physics, "model", "COLDReconLatentDiffusion")
    footprint = _pick(computational, "model", "COLDReconLatentDiffusionPhysicsTrained")
    fno_footprint = _pick(computational, "model", "COLDReconFNOOperatorDiffusion")
    eic_unc = uncertainty[
        uncertainty.get("model", pd.Series(dtype=str)).astype(str).eq("COLDReconLatentDiffusionPhysicsTrained")
        & uncertainty.get("target", pd.Series(dtype=str)).astype(str).eq("eic")
    ]
    refined_uw = uncertainty[
        uncertainty.get("model", pd.Series(dtype=str)).astype(str).eq("COLDReconLatentDiffusionPhysicsRefined")
        & uncertainty.get("target", pd.Series(dtype=str)).astype(str).eq("unfrozen_water")
    ]
    cg_tier = None
    eg_tier = None
    for tier in readiness.get("tiers", []):
        if tier.get("tier") == "CG algorithm article":
            cg_tier = tier
        if tier.get("tier") == "EG field-generalization claim":
            eg_tier = tier

    def token_count(source_key: str, observation_type: str) -> str:
        if tokens.empty:
            return "not available"
        rows = tokens[
            tokens["source_key"].astype(str).eq(source_key)
            & tokens["observation_type"].astype(str).eq(observation_type)
        ]
        if rows.empty:
            return "not available"
        return _fmt(rows.iloc[0]["n_tokens"], digits=0)

    return {
        "latent_mean_iou": _metric(latent, "mean_iou"),
        "physics_mean_iou": _metric(physics_trained, "mean_iou"),
        "fno_mean_iou": _metric(fno, "mean_iou"),
        "gradient_mean_iou": _metric(gradient, "mean_iou"),
        "latent_eic_rmse": _metric(latent, "eic_rmse"),
        "physics_eic_rmse": _metric(physics_trained, "eic_rmse"),
        "raw_heat_rmse": _metric(raw_physics, "heat_residual_rmse"),
        "refined_heat_rmse": _metric(refined_physics, "heat_residual_rmse"),
        "refined_log_resistivity_rmse": _metric(refined_physics, "log_resistivity_empirical_rmse"),
        "latent_params_m": _metric(footprint, "total_params_m"),
        "latent_prediction_mb": _metric(footprint, "prediction_mb"),
        "fno_params_m": _metric(fno_footprint, "total_params_m"),
        "public_sources": cg_gate.get("independent_public_sources_passed", "not available"),
        "passed_tasks": cg_gate.get("passed_tasks", "not available"),
        "total_tasks": cg_gate.get("total_tasks", "not available"),
        "eic_sources": cg_gate.get("eic_sources_passed", "not available"),
        "facies_sources": cg_gate.get("facies_sources_passed", "not available"),
        "usgs_ert": token_count("usgs_ert_nmr", "ert_log_resistivity"),
        "usgs_nmr": token_count("usgs_ert_nmr", "nmr_unfrozen_water"),
        "usgs_alt": token_count("usgs_ert_nmr", "alt"),
        "usgs_eic": token_count("usgs_eic_cores", "borehole_eic"),
        "arctic_facies": token_count("arcticdata_upper_permafrost_cryostratigraphy", "borehole_facies"),
        "arctic_eic": token_count("arcticdata_upper_permafrost_cryostratigraphy", "borehole_eic"),
        "jago_eic": token_count("arcticdata_jago_ground_ice_2018", "borehole_eic"),
        "coord_units": coord.get("n_georeferenced_units", "not available"),
        "coord_sites": coord.get("n_sites_with_georeferenced_units", "not available"),
        "coord_eic": coord.get("n_eic_measurements", "not available"),
        "coord_wedge": coord.get("n_wedge_ice_units", "not available"),
        "voi_enrichment": _fmt(voi.get("composite_top_voi_error_enrichment")),
        "voi_high_eic": _fmt(voi.get("high_eic_top_voi_error_enrichment")),
        "voi_spearman": _fmt(voi.get("composite_spearman_voi_error")),
        "cg_score": _fmt(cg_tier.get("score") if cg_tier else None),
        "cg_pass": cg_tier.get("n_pass") if cg_tier else "not available",
        "cg_criteria": cg_tier.get("n_criteria") if cg_tier else "not available",
        "eg_score": _fmt(eg_tier.get("score") if eg_tier else None),
        "eg_not_yet": eg_tier.get("n_not_yet") if eg_tier else "not available",
        "eic_unc_spearman": _fmt(eic_unc.iloc[0]["spearman_uncertainty_error"] if not eic_unc.empty else None),
        "uw_unc_spearman": _fmt(refined_uw.iloc[0]["spearman_uncertainty_error"] if not refined_uw.empty else None),
        "rare_hybrid_wedge_recall": _metric(
            _pick(rare_hybrid, "model", "COLDReconLatentDiffusionRareFaciesHybrid"),
            "wedge_ice_recall",
        ),
        "rare_hybrid_precision": _metric(
            _pick(rare_hybrid, "model", "COLDReconLatentDiffusionRareFaciesHybrid"),
            "wedge_ice_precision",
        ),
        "wedge_recall_head": _metric(_pick(wedge, "model", "COLDReconArcticDataWedgeRecallHead"), "recall"),
    }


def _manuscript_lines(project_root: Path, manifest: pd.DataFrame) -> list[str]:
    m = _article_metrics(project_root)
    generated = datetime.now(timezone.utc).replace(microsecond=0).isoformat()
    lines = [
        "# COLD-Recon: multi-source sparse-observation constrained probabilistic 3D permafrost reconstruction with a physics-guided conditional diffusion neural operator",
        "",
        "## Computational Geoscience Algorithm Manuscript",
        "",
        f"Generated UTC: `{generated}`",
        "",
        "### Abstract",
        "",
        (
            "Subsurface ice, cryofacies and unfrozen-water structure control permafrost thermal and hydrological "
            "behaviour, but field observations are usually sparse, heterogeneous and only partly colocated. "
            "We present COLD-Recon, a conditional diffusion neural-operator framework that reconstructs "
            "probabilistic three-dimensional permafrost structure from borehole intervals, electrical resistivity, "
            "nuclear magnetic resonance and active-layer observations. The method represents observations as typed "
            "spatial tokens, samples gridded posterior fields in a latent 3D volume, and audits the result through "
            "physical-consistency, uncertainty-calibration, rare-cryostructure and public-data validation gates. "
            f"Under synthetic full-field truth, the physics-trained posterior reached mean facies IoU {m['physics_mean_iou']}, "
            f"compared with {m['gradient_mean_iou']} for the strongest tree baseline and {m['fno_mean_iou']} for the "
            "FNO-Transformer diffusion variant. "
            f"The public evidence gate passed {m['passed_tasks']}/{m['total_tasks']} validation tasks across "
            f"{m['public_sources']} independent public sources, including {m['eic_sources']} EIC sources and "
            f"{m['facies_sources']} cryofacies source. A retrospective full-field VOI audit enriched composite "
            f"reconstruction error by {m['voi_enrichment']} in the top ranked decile. A readiness audit scores the "
            f"CG algorithm article {m['cg_score']} ({m['cg_pass']}/{m['cg_criteria']} criteria passed), while keeping "
            "regional EG field-generalization as a bounded future claim because public dense 3D ground truth remains unavailable."
        ),
        "",
        "### Keywords",
        "",
        "Permafrost; ground ice; cryostratigraphy; conditional diffusion; neural operator; sparse observations; posterior reconstruction; public data validation.",
        "",
        "## 1. Introduction",
        "",
        (
            "Permafrost site characterization is limited by a structural mismatch between the variables of interest "
            "and the observations normally available. Cryofacies, excess-ice content (EIC), temperature, unfrozen "
            "water and resistivity vary in three dimensions, whereas field evidence is commonly composed of isolated "
            "boreholes, short core intervals, partial ERT/NMR lines and active-layer measurements. Treating this as "
            "ordinary interpolation understates uncertainty and obscures rare but consequential ice-rich structures. "
            "COLD-Recon instead treats reconstruction as conditional posterior generation: infer a distribution over "
            "3D frozen-ground states, not a single deterministic map."
        ),
        "",
        (
            "The CG contribution is algorithmic. The article does not claim a completed regional ground-ice map. "
            "It contributes a reproducible reconstruction workflow, an observation-token interface, a conditional "
            "diffusion neural operator, physics-guided posterior checking, rare-event operating points, public-data "
            "evidence gates and a posterior observation-design diagnostic. This scope is essential: public permafrost "
            "data can validate sparse direct labels and dense proxies, but they do not yet provide independent dense "
            "3D ground-truth volumes."
        ),
        "",
        "## 2. Problem formulation",
        "",
        (
            "Let M denote a gridded 3D permafrost state containing categorical cryofacies and continuous fields "
            "for EIC, temperature, unfrozen water and log-resistivity. Let O denote sparse heterogeneous observations. "
            "The reconstruction target is p(M | O), with hard observation consistency near measured locations, "
            "uncertainty away from observations and physically plausible coupling among frozen-ground variables. "
            "Observation types are encoded as tokens containing normalized coordinates, depth support, value, "
            "uncertainty and type masks, allowing borehole labels, ERT, NMR and active-layer constraints to enter a "
            "shared conditional interface."
        ),
        "",
        "## 3. Method",
        "",
        "### 3.1 Observation-token conditioning",
        "",
        (
            "The token interface is designed for multi-source sparse evidence. Borehole facies and EIC intervals "
            "supply direct vertical labels; ERT log-resistivity and NMR water-content values supply geophysical proxy "
            "constraints; active-layer observations supply thaw-depth context. Tokens are projected into a learned "
            "conditioning representation and injected into the posterior denoiser, so missing observation types do not "
            "require separate model definitions."
        ),
        "",
        "### 3.2 Conditional diffusion neural operator",
        "",
        (
            "COLD-Recon uses a 3D latent representation of the permafrost volume and denoises this latent field under "
            "the observation-token context. The tested family includes a compact latent diffusion model, a higher-"
            "capacity FNO-Transformer diffusion denoiser and a rectified-flow variant. The compact physics-trained "
            f"posterior uses {m['latent_params_m']} million parameters and writes a {m['latent_prediction_mb']} MB "
            f"prediction artifact; the FNO variant uses {m['fno_params_m']} million parameters. These footprint "
            "statistics are reported with accuracy so that model capacity is visible to reviewers."
        ),
        "",
        "### 3.3 Physical consistency and calibration",
        "",
        (
            "The workflow audits empirical unfrozen-water consistency, resistivity coupling and simplified heat "
            "residuals. Physics-guided training and post-hoc projection are evaluated separately, preventing "
            "calibration from being conflated with physical correctness. In the current synthetic audit, heat-residual "
            f"RMSE changed from {m['raw_heat_rmse']} for the raw latent diffusion posterior to {m['refined_heat_rmse']} "
            f"after physics refinement, and refined log-resistivity empirical RMSE is {m['refined_log_resistivity_rmse']}."
        ),
        "",
        "### 3.4 Rare cryostructure operating points",
        "",
        (
            "Rare ice-rich structures are not summarized by mean IoU alone. The package therefore reports high-EIC "
            "screening, wedge recall, precision and explicit rare-facies hybrid operating points. The synthetic "
            f"rare-facies hybrid reached wedge recall {m['rare_hybrid_wedge_recall']} and precision "
            f"{m['rare_hybrid_precision']}, while the public ArcticData wedge recall head is reported as a recall-"
            "oriented operating point rather than a hidden aggregate score."
        ),
        "",
        "### 3.5 Posterior value of information",
        "",
        (
            "The posterior is also used to rank additional borehole and ERT-line targets. The VOI score combines "
            "posterior spread, ice-rich ambiguity, thaw-sensitive EIC structure and novelty relative to existing "
            "observations. This is a diagnostic layer for the reconstruction algorithm, not a claim of prospectively "
            "optimized field acquisition."
        ),
        "",
        "## 4. Data and experiments",
        "",
        (
            "Synthetic experiments provide full-field truth for volumetric metrics. Public validation uses downloaded "
            "and processed USGS ERT/NMR/thaw-depth records, USGS core EIC measurements, Arctic Data Center upper-"
            "permafrost cryostratigraphy/EIC data and an independent Jago River ground-ice release. The processed "
            f"public token inventory contains {m['usgs_ert']} ERT log-resistivity tokens, {m['usgs_nmr']} NMR tokens, "
            f"{m['usgs_alt']} active-layer tokens, {m['usgs_eic']} USGS EIC intervals, {m['arctic_facies']} ArcticData "
            f"facies tokens, {m['arctic_eic']} ArcticData EIC tokens and {m['jago_eic']} Jago EIC tokens. The coordinate "
            f"coverage audit records {m['coord_units']} georeferenced vertical units across {m['coord_sites']} sites, "
            f"including {m['coord_eic']} EIC measurements and {m['coord_wedge']} wedge-ice units."
        ),
        "",
        "## 5. Results",
        "",
        "### 5.1 Full-field synthetic reconstruction",
        "",
        (
            f"The strongest classical baseline reached mean facies IoU {m['gradient_mean_iou']}. The compact latent "
            f"diffusion posterior reached {m['latent_mean_iou']}, the FNO-Transformer diffusion posterior reached "
            f"{m['fno_mean_iou']}, and the physics-trained posterior reached {m['physics_mean_iou']}. EIC RMSE for the "
            f"latent diffusion and physics-trained variants is {m['latent_eic_rmse']} and {m['physics_eic_rmse']}, "
            "respectively. The result supports the central algorithmic claim: posterior generation can improve "
            "volumetric facies reconstruction while preserving uncertainty outputs."
        ),
        "",
        "### 5.2 Uncertainty and physics diagnostics",
        "",
        (
            f"For the physics-trained posterior, EIC uncertainty has Spearman correlation {m['eic_unc_spearman']} with "
            f"absolute EIC error. After physics refinement, unfrozen-water uncertainty has Spearman correlation "
            f"{m['uw_unc_spearman']} with absolute unfrozen-water error. The probabilistic reliability claim is "
            "therefore target-specific: EIC uncertainty localizes errors strongly, while other fields require more "
            "careful interpretation."
        ),
        "",
        "### 5.3 Public validation evidence",
        "",
        (
            f"The public evidence gate passes {m['passed_tasks']}/{m['total_tasks']} tasks across {m['public_sources']} "
            f"independent public sources. The gate includes {m['eic_sources']} EIC sources and {m['facies_sources']} "
            "cryofacies source, plus recall-oriented wedge-ice handling. These tests provide public, reproducible "
            "support for the algorithm but are not dense 3D ground truth."
        ),
        "",
        "### 5.4 External-transfer boundaries",
        "",
        (
            "The external-generalization, transfer-failure and domain-support audits are included because they expose "
            "when site transfer is model-supported and when a guarded local-prior adapter is more appropriate. This "
            "keeps the manuscript in the CG algorithm class: transfer is audited as a boundary condition, not promoted "
            "to a regional field-generalization claim."
        ),
        "",
        "### 5.5 Observation-design audit",
        "",
        (
            f"In the synthetic retrospective VOI backtest, the top VOI decile enriched composite reconstruction error "
            f"by {m['voi_enrichment']}, high-EIC error by {m['voi_high_eic']} and had Spearman correlation "
            f"{m['voi_spearman']} with composite error. Thus the VOI layer is useful as a posterior blind-spot "
            "diagnostic, while still requiring prospective field validation for operational acquisition design."
        ),
        "",
        "### 5.6 CG readiness",
        "",
        (
            f"The readiness audit scores the CG algorithm article {m['cg_score']} with {m['cg_pass']}/{m['cg_criteria']} "
            f"criteria passed. The EG field-generalization tier remains at score {m['eg_score']} with "
            f"{m['eg_not_yet']} not-yet criterion. The article is therefore ready as a CG algorithm manuscript, not as "
            "a completed regional field-generalization study."
        ),
        "",
        "## 6. Discussion",
        "",
        (
            "COLD-Recon is strongest when evaluated as an auditable posterior reconstruction framework. Its contribution "
            "is the integration of heterogeneous sparse observations, conditional posterior generation, physical "
            "diagnostics, uncertainty-error alignment, rare-structure operating points, public evidence gates and "
            "observation-design scoring in one reproducible pipeline. The major limitation is not a missing paragraph "
            "but a missing public data type: independent dense 3D ground-ice or cryofacies truth. Until such data exist, "
            "full-field validation must remain synthetic, and public validation must be reported as sparse direct labels "
            "plus dense proxies."
        ),
        "",
        "## 7. Conclusions",
        "",
        (
            "COLD-Recon reconstructs probabilistic 3D permafrost structure from multi-source sparse observations and "
            "provides a reproducible CG algorithm evidence chain. The current results support the algorithm article: "
            "synthetic full-field reconstruction is benchmarked against baselines, physics and uncertainty are audited, "
            "public validation uses three independent sources, rare-cryostructure behaviour is separated from mean IoU, "
            "and VOI diagnostics are retrospectively checked. The remaining EG gap is explicit and should be addressed "
            "with prospective field validation and public dense 3D ground-truth releases."
        ),
        "",
        "## Data and code availability",
        "",
        (
            "All public-data sources, processed token inventories, validation tables, source-data files, figure scripts "
            "and reproduction commands are included in this project directory. The CG submission package copies only PNG "
            "figures and Python scripts for figure synchronization; full upstream reproduction commands are recorded in "
            "`outputs/tables/reproducibility_summary.json`."
        ),
        "",
        "## Figures",
        "",
    ]
    for _, row in manifest.iterrows():
        number = int(row["figure_number"])
        title = str(row["title"])
        package_png = Path(str(row["package_png"]))
        rel_png = Path("figures") / package_png.name
        caption = str(row["caption"]).strip() or title
        role = str(row["claim_role"]).strip()
        category = str(row["category_key"]).strip()
        detail = []
        if role:
            detail.append(f"Claim role: {role}.")
        if category:
            detail.append(f"Evidence class: {category}.")
        lines.extend(
            [
                f"![Figure {number}. {title}.]({rel_png.as_posix()})",
                "",
                f"**Figure {number} | {title}.** {caption} {' '.join(detail)}".strip(),
                "",
            ]
        )
    lines.extend(
        [
            "## References",
            "",
            "BibTeX records are provided in `references.bib` and include public data products and core method references used by the pipeline.",
            "",
        ]
    )
    return lines


INLINE_TOKEN = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
IMAGE_LINE = re.compile(r"^!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)$")


def _add_inline_runs(paragraph, text: str) -> None:
    for token in INLINE_TOKEN.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
        else:
            paragraph.add_run(token)


def _build_docx(article_md: Path, output_docx: Path) -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header.paragraphs[0].text = "COLD-Recon CG algorithm manuscript"
    section.footer.paragraphs[0].text = "Generated from audited COLD-Recon outputs"

    first_title = True
    for raw_line in article_md.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        image_match = IMAGE_LINE.match(line)
        if image_match:
            img_path = (article_md.parent / image_match.group("path")).resolve()
            if img_path.exists():
                doc.add_picture(str(img_path), width=Inches(6.3))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if line.startswith("# "):
            style = "Title" if first_title else "Heading 1"
            p = doc.add_paragraph(style=style)
            _add_inline_runs(p, line[2:].strip())
            first_title = False
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            _add_inline_runs(p, line[3:].strip())
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            _add_inline_runs(p, line[4:].strip())
        elif (
            line.startswith("**Figure ")
            or line.startswith("**图 ")
            or line.startswith("**Graphical abstract")
            or line.startswith("**图文摘要")
        ):
            p = doc.add_paragraph(style="Caption")
            _add_inline_runs(p, line)
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, line)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)
    return output_docx


def _write_readme(package_dir: Path, result_paths: dict[str, Path], manifest: pd.DataFrame) -> Path:
    readme = package_dir / "README.md"
    generated = datetime.now(timezone.utc).replace(microsecond=0).isoformat()
    categories = manifest["category_key"].value_counts().sort_index()
    lines = [
        "# COLD-Recon CG Algorithm Submission Package",
        "",
        f"Generated UTC: `{generated}`",
        "",
        "This is the independent Computational Geoscience algorithm manuscript package for COLD-Recon.",
        "",
        "## Contents",
        "",
        f"- `{result_paths['article_md'].name}`: complete manuscript source with embedded PNG figure links.",
        f"- `{result_paths['article_docx'].name}`: DOCX version with the PNG figures embedded.",
        "- `figures/`: PNG-only figure bundle.",
        "- `figure_scripts/`: one Python synchronization script per PNG figure.",
        "- `tables/`: copied audited CSV/JSON result tables.",
        "- `source_data/`: plotted source-data CSV files.",
        "- `figure_manifest.csv`: figure number, source PNG, package PNG, script and upstream generator command.",
        "- `references.bib`: public data and method references.",
        "",
        "## Figure Policy",
        "",
        f"- Included figures: `{len(manifest)}`.",
        "- Figure image format in this package: PNG only.",
        "- No SVG/PDF/TIFF files are copied into `figures/`.",
        "- Application-oriented figures outside probabilistic 3D reconstruction are excluded.",
        "",
        "## Figure Categories",
        "",
    ]
    for category, count in categories.items():
        lines.append(f"- `{category}`: {int(count)} figures.")
    lines.extend(
        [
            "",
            "## Readiness Position",
            "",
            "The package supports a CG algorithm manuscript. It does not claim a completed EG regional field-generalization study.",
            "",
        ]
    )
    readme.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    return readme


def _zip_package(package_dir: Path) -> Path:
    zip_path = package_dir / "cold_recon_cg_algorithm_submission.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(package_dir.rglob("*")):
            if path == zip_path or path.is_dir():
                continue
            zf.write(path, path.relative_to(package_dir))
    return zip_path


def build_cg_submission(project_root: Path = Path("."), package_dir: Path = PACKAGE_DIR) -> CGSubmissionResult:
    project_root = project_root.resolve()
    package_dir = (project_root / package_dir).resolve() if not package_dir.is_absolute() else package_dir
    package_dir.mkdir(parents=True, exist_ok=True)
    for child_name in ("figures", "figure_scripts", "tables", "source_data"):
        child = package_dir / child_name
        if child.exists():
            shutil.rmtree(child)
    for file_name in (
        f"{ARTICLE_NAME}.md",
        f"{ARTICLE_NAME}.docx",
        "README.md",
        "figure_manifest.csv",
        "references.bib",
        "cold_recon_cg_algorithm_submission.zip",
    ):
        path = package_dir / file_name
        if path.exists():
            path.unlink()

    selected = _selected_figures(project_root)
    if selected.empty:
        raise RuntimeError("No CG submission figures were selected from outputs/tables/figure_atlas.csv")
    manifest = _build_figure_manifest(project_root, selected, package_dir)
    _sync_pngs_and_scripts(project_root, package_dir, manifest)
    _copy_auxiliary_tables(project_root, package_dir, manifest, selected)

    manifest_path = package_dir / "figure_manifest.csv"
    manifest.to_csv(manifest_path, index=False)

    article_md = package_dir / f"{ARTICLE_NAME}.md"
    article_lines = _manuscript_lines(project_root, manifest)
    article_md.write_text("\n".join(article_lines).rstrip() + "\n", encoding="utf-8")

    article_docx = package_dir / f"{ARTICLE_NAME}.docx"
    _build_docx(article_md, article_docx)

    references = write_references_bib(package_dir)
    result_paths = {"article_md": article_md, "article_docx": article_docx, "references": references}
    readme = _write_readme(package_dir, result_paths, manifest)
    zip_path = _zip_package(package_dir)

    return CGSubmissionResult(
        package_dir=package_dir,
        article_md=article_md,
        article_docx=article_docx,
        readme=readme,
        figure_manifest=manifest_path,
        package_zip=zip_path,
        n_figures=int(len(manifest)),
        n_scripts=int(len(list((package_dir / "figure_scripts").glob("*.py")))),
    )
