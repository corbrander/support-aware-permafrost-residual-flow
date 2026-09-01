from __future__ import annotations

import json
import re
import shutil
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


FIGURE_STEMS = (
    "nature_figure_1_overview",
    "nature_figure_2_real_data_gate",
    "nature_figure_3_cited_ground_ice",
    "nature_figure_4_site_investigation",
    "innovation_positioning_audit",
    "arcticdata_wedge_operating_curve",
    "synthetic_rare_cryostructure_audit",
    "rare_facies_hybrid_operating_curve",
    "posterior_uncertainty_alignment",
    "computational_footprint_summary",
    "voi_backtest_audit",
    "coordinate_label_coverage_audit",
    "external_generalization_audit",
    "transfer_failure_attribution",
    "domain_support_audit",
    "journal_readiness_audit",
)
FIGURE_EXTENSIONS = (".svg", ".pdf", ".png", ".tiff")
SOURCE_DATA_FILES = (
    "nature_figure_1_source_data.csv",
    "nature_figure_2_source_data.csv",
    "nature_figure_3_source_data.csv",
    "nature_figure_4_source_data.csv",
    "innovation_positioning_audit_source_data.csv",
    "voi_backtest_audit_source_data.csv",
    "coordinate_label_coverage_audit_source_data.csv",
    "external_generalization_audit_source_data.csv",
    "transfer_failure_attribution_source_data.csv",
    "domain_support_audit_source_data.csv",
    "journal_readiness_audit_source_data.csv",
)
TABLE_FILES = (
    "figure_atlas.csv",
    "innovation_positioning_audit.csv",
    "innovation_positioning_summary.json",
    "voi_backtest_audit.csv",
    "voi_backtest_summary.json",
    "coordinate_label_coverage_audit.csv",
    "coordinate_label_coverage_summary.json",
    "domain_support_site_audit.csv",
    "domain_support_summary.json",
    "journal_readiness_audit.csv",
    "journal_readiness_summary.json",
    "real_data_cg_gate.json",
    "real_data_cg_benchmark.csv",
    "public_data_token_inventory.csv",
    "public_data_provenance.csv",
    "site_investigation_boreholes.csv",
    "site_investigation_ert_lines.csv",
    "arcticdata_wedge_probability_holdout_scores.csv",
    "arcticdata_wedge_operating_curve.csv",
    "arcticdata_wedge_operating_points.csv",
    "synthetic_rare_cryostructure_audit.csv",
    "diffusion_rare_facies_hybrid_metrics.csv",
    "rare_facies_hybrid_operating_curve.csv",
    "posterior_uncertainty_alignment.csv",
    "computational_footprint.csv",
    "external_generalization_site_deltas.csv",
    "external_generalization_audit.csv",
    "transfer_failure_site_diagnostics.csv",
    "transfer_failure_attribution_summary.csv",
)

PACKAGE_COLUMN_RENAMES = {
    "settlement_risk": "thaw_sensitive_eic_proxy",
    "differential_settlement": "eic_gradient_proxy",
    "settlement_potential": "thaw_sensitive_eic_raw",
    "weighted_settlement_risk": "weighted_thaw_sensitive_eic_proxy",
    "weighted_differential_settlement": "weighted_eic_gradient_proxy",
}
PACKAGE_VALUE_RENAMES = {
    "settlement_risk": "thaw_sensitive_eic_proxy",
    "differential_settlement": "eic_gradient_proxy",
}


@dataclass(frozen=True)
class SubmissionPackageResult:
    article_docx: Path
    package_dir: Path
    package_readme: Path
    package_zip: Path


def _set_font(font, name: str, size_pt: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    font.name = name
    if size_pt is not None:
        font.size = Pt(size_pt)
    if color is not None:
        font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    if bold is not None:
        font.bold = bold


def _set_style_font(style, name: str, size_pt: float, color: str | None = None, bold: bool | None = None) -> None:
    _set_font(style.font, name, size_pt, color, bold)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    _set_style_font(normal, "Calibri", 11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    title = styles["Title"]
    _set_style_font(title, "Calibri", 18, "0B2545", True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = styles["Subtitle"]
    _set_style_font(subtitle, "Calibri", 11, "555555", False)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle.paragraph_format.line_spacing = 1.167

    h1 = styles["Heading 1"]
    _set_style_font(h1, "Calibri", 16, "2E74B5", True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    _set_style_font(h2, "Calibri", 13, "2E74B5", True)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    _set_style_font(h3, "Calibri", 12, "1F4D78", True)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    for style_name in ("List Number", "List Bullet"):
        style = styles[style_name]
        _set_style_font(style, "Calibri", 11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    if "Figure Caption" not in styles:
        caption = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Figure Caption"]
    _set_style_font(caption, "Calibri", 9, "333333")
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(12)
    caption.paragraph_format.line_spacing = 1.167

    if "Lead" not in styles:
        lead = styles.add_style("Lead", WD_STYLE_TYPE.PARAGRAPH)
    else:
        lead = styles["Lead"]
    _set_style_font(lead, "Calibri", 11, "0B2545")
    lead.paragraph_format.space_before = Pt(0)
    lead.paragraph_format.space_after = Pt(10)
    lead.paragraph_format.line_spacing = 1.25


def _configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header_p = section.header.paragraphs[0]
    header_p.text = "COLD-Recon Nature-style submission draft"
    header_p.style = doc.styles["Subtitle"]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer_p = section.footer.paragraphs[0]
    footer_p.text = "Generated from audited COLD-Recon reproducibility outputs"
    footer_p.style = doc.styles["Subtitle"]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _set_core_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "Multi-source sparse-observation constrained probabilistic 3D permafrost reconstruction with a physics-guided conditional diffusion neural operator"
    props.subject = "Complete Nature-style submission draft and reproducible package"
    props.author = "COLD-Recon reproducibility pipeline"
    props.comments = "Generated by cold_recon.scripts.48_make_submission_package"


INLINE_TOKEN = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
IMAGE_LINE = re.compile(r"^!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)$")


def _add_inline_runs(paragraph, text: str) -> None:
    for token in INLINE_TOKEN.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
            run.font.size = Pt(9)
        else:
            paragraph.add_run(token)


def _add_paragraph(doc: Document, text: str, style: str = "Normal") -> None:
    p = doc.add_paragraph(style=style)
    _add_inline_runs(p, text)


def _resolve_asset(article_path: Path, raw_path: str) -> Path:
    path = Path(raw_path)
    if path.is_absolute():
        return path
    return (article_path.parent / path).resolve()


def _add_image(doc: Document, image_path: Path) -> None:
    if not image_path.exists():
        _add_paragraph(doc, f"[Missing figure asset: {image_path}]", style="Figure Caption")
        return
    doc.add_picture(str(image_path), width=Inches(6.35))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_article_docx(article_md: Path, output_docx: Path) -> Path:
    doc = Document()
    _configure_page(doc)
    _configure_styles(doc)
    _set_core_properties(doc)

    in_figure_captions = False
    first_title = True
    figure_count = 0
    for raw_line in article_md.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue

        image_match = IMAGE_LINE.match(line)
        if image_match:
            figure_count += 1
            if in_figure_captions:
                doc.add_page_break()
            _add_image(doc, _resolve_asset(article_md, image_match.group("path")))
            continue

        if line == "### Figure Captions":
            in_figure_captions = True
            _add_paragraph(doc, "Figure Captions", style="Heading 1")
            continue

        if line.startswith("# "):
            p = doc.add_paragraph(style="Title" if first_title else "Heading 1")
            _add_inline_runs(p, line[2:].strip())
            first_title = False
        elif line.startswith("## "):
            text = line[3:].strip()
            style = "Subtitle" if text in {"Nature-style Article Draft", "Complete Submission Draft"} else "Heading 1"
            _add_paragraph(doc, text, style=style)
        elif line.startswith("### "):
            _add_paragraph(doc, line[4:].strip(), style="Heading 1")
        elif line.startswith("#### "):
            _add_paragraph(doc, line[5:].strip(), style="Heading 2")
        elif re.match(r"^\d+\. ", line):
            _add_paragraph(doc, re.sub(r"^\d+\. ", "", line), style="List Number")
        elif line.startswith("- "):
            _add_paragraph(doc, line[2:].strip(), style="List Bullet")
        elif line.startswith("**One-sentence argument.**"):
            _add_paragraph(doc, line, style="Lead")
        elif line.startswith("**Figure "):
            _add_paragraph(doc, line, style="Figure Caption")
        else:
            _add_paragraph(doc, line)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)
    return output_docx


def _copy_file(src: Path, dst: Path) -> Path:
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)
    return dst


def _copy_csv_with_algorithm_labels(src: Path, dst: Path) -> Path:
    dst.parent.mkdir(parents=True, exist_ok=True)
    df = pd.read_csv(src)
    df = df.rename(columns=PACKAGE_COLUMN_RENAMES)
    if "component" in df.columns:
        df["component"] = df["component"].map(lambda value: PACKAGE_VALUE_RENAMES.get(str(value), value))
    df.to_csv(dst, index=False)
    return dst


def _clean_submission_text(text: object) -> str:
    cleaned = str(text).strip()
    replacements = {
        "supplemental observation": "additional observation",
        "supplemental boreholes": "additional boreholes",
        "supplemental ERT": "additional ERT",
        "Supplemental observation": "Additional observation",
        "Supplemental boreholes": "Additional boreholes",
        "Supplemental ERT": "Additional ERT",
    }
    for old, new in replacements.items():
        cleaned = cleaned.replace(old, new)
    return cleaned


def _truthy(value: object) -> bool:
    return str(value).strip().lower() in {"1", "true", "yes", "y"}


def _submission_figure_atlas(atlas: pd.DataFrame) -> pd.DataFrame:
    if atlas.empty or "copy_to_submission" not in atlas.columns:
        return atlas.iloc[0:0].copy()
    filtered = atlas[atlas["copy_to_submission"].map(_truthy)].copy()
    if "manuscript_status" in filtered.columns:
        filtered = filtered[~filtered["manuscript_status"].astype(str).eq("scope_boundary_excluded")]
    sort_columns = [col for col in ["category_order", "category_key", "stem"] if col in filtered.columns]
    if sort_columns:
        filtered = filtered.sort_values(sort_columns, kind="stable")
    return filtered


def _copy_submission_figure_atlas_csv(src: Path, dst: Path) -> Path:
    dst.parent.mkdir(parents=True, exist_ok=True)
    atlas = _submission_figure_atlas(pd.read_csv(src))
    for column in ("caption", "boundary_note"):
        if column in atlas.columns:
            atlas[column] = atlas[column].map(_clean_submission_text)
    atlas.to_csv(dst, index=False)
    return dst


def _render_submission_figure_atlas(atlas: pd.DataFrame) -> str:
    lines = [
        "# Complete Submission Figure Atlas",
        "",
        "This package-level atlas lists the algorithm, public-data, physics, uncertainty, rare-cryostructure and observation-design figures copied into the submission package and embedded in the complete submission draft.",
        "",
    ]
    if atlas.empty:
        lines.append("No atlas-selected submission figures were found.")
        return "\n".join(lines)
    for _, row in atlas.iterrows():
        stem = str(row.get("stem", "")).strip()
        preferred = str(row.get("preferred_path", "")).strip()
        caption = _clean_submission_text(row.get("caption", ""))
        category = str(row.get("category", "")).strip()
        claim_role = str(row.get("claim_role", "")).strip()
        if not stem:
            continue
        lines.extend(
            [
                f"## {stem}",
                "",
                f"- Category: {category or 'algorithm evidence'}",
                f"- Claim role: {claim_role or 'submission evidence'}",
                f"- Caption: {caption or stem}",
            ]
        )
        if preferred:
            lines.extend(["", f"![{stem}](supplementary_figures/{Path(preferred).name})"])
        lines.append("")
    return "\n".join(lines)


def _copy_supplementary_figure_atlas(root: Path, package_dir: Path) -> None:
    paper_dir = root / "paper"
    table_dir = root / "outputs" / "tables"
    atlas_csv = table_dir / "figure_atlas.csv"
    if not atlas_csv.exists():
        return
    atlas = _submission_figure_atlas(pd.read_csv(atlas_csv))
    for column in ("caption", "boundary_note"):
        if column in atlas.columns:
            atlas[column] = atlas[column].map(_clean_submission_text)
    (package_dir / "tables").mkdir(parents=True, exist_ok=True)
    atlas.to_csv(package_dir / "tables" / atlas_csv.name, index=False)
    (package_dir / "supplementary_figure_atlas.md").write_text(_render_submission_figure_atlas(atlas), encoding="utf-8")
    if "all_paths" not in atlas.columns:
        return
    for _, row in atlas.iterrows():
        for rel_path in str(row["all_paths"]).split(";"):
            rel_path = rel_path.strip()
            if not rel_path:
                continue
            src = root / rel_path
            if src.exists():
                _copy_file(src, package_dir / "supplementary_figures" / src.name)


def _copy_all_submission_figures(root: Path, package_dir: Path) -> int:
    atlas_csv = root / "outputs" / "tables" / "figure_atlas.csv"
    if not atlas_csv.exists():
        return 0
    atlas = _submission_figure_atlas(pd.read_csv(atlas_csv))
    if "all_paths" not in atlas.columns:
        return 0
    copied_stems: set[str] = set()
    for _, row in atlas.iterrows():
        stem = str(row.get("stem", "")).strip()
        if stem:
            copied_stems.add(stem)
        for rel_path in str(row["all_paths"]).split(";"):
            rel_path = rel_path.strip()
            if not rel_path:
                continue
            src = root / rel_path
            if src.exists():
                _copy_file(src, package_dir / "figures" / src.name)
    return len(copied_stems)


def _count_submission_figures(root: Path) -> int:
    atlas_csv = root / "outputs" / "tables" / "figure_atlas.csv"
    if not atlas_csv.exists():
        return len(FIGURE_STEMS)
    return int(len(_submission_figure_atlas(pd.read_csv(atlas_csv))))


def _iter_package_files(package_dir: Path, package_zip: Path) -> Iterable[Path]:
    for path in sorted(package_dir.rglob("*")):
        if path.is_file() and path.resolve() != package_zip.resolve():
            yield path


def _write_package_readme(root: Path, package_dir: Path) -> Path:
    table_dir = root / "outputs" / "tables"
    gate_path = table_dir / "real_data_cg_gate.json"
    gate = json.loads(gate_path.read_text(encoding="utf-8")) if gate_path.exists() else {}
    submission_figure_count = _count_submission_figures(root)
    lines = [
        "# COLD-Recon Submission Package",
        "",
        f"Generated UTC: `{datetime.now(timezone.utc).isoformat()}`",
        "",
        "## Primary Manuscript Files",
        "",
        f"- `cold_recon_nature_article.docx`: complete submission draft with the full atlas-selected main-manuscript figure set embedded (currently {submission_figure_count} figure stems), using the `narrative_proposal` DOCX preset.",
        "- `cold_recon_nature_article.md`: source Markdown for the complete submission text.",
        "- `cold_recon_claim_evidence_audit.md`: claim-evidence map, reviewer boundary audit, and CG/EG readiness boundary.",
        "- `supplementary_figure_atlas.md`: package-level atlas of the algorithm/public-data figures embedded in the complete submission draft.",
        "- `references.bib`: BibTeX records for public data products and method references.",
        "",
        "## Comprehensive Submission Figure Bundle",
        "",
        f"- All `{submission_figure_count}` atlas-selected algorithm/public-data figure stems are embedded in the complete submission draft and copied to `figures/`.",
        "- `figures/nature_figure_1_overview.*`: Figure 1, workflow, synthetic benchmark, calibration and physics diagnostics.",
        "- `figures/nature_figure_2_real_data_gate.*`: Figure 2, public evidence gate and pass/fail matrix.",
        "- `figures/nature_figure_3_cited_ground_ice.*`: Figure 3, cited ArcticData, USGS and Jago River ground-ice validation data.",
        "- `figures/nature_figure_4_site_investigation.*`: Figure 4, posterior value-of-information diagnostic, borehole targets and ERT-line targets.",
        "- `figures/innovation_positioning_audit.*`: Figure 5, evidence-mapped novelty and positioning audit.",
        "- `figures/external_generalization_audit.*`: Figure 6, public multi-site generalization and boundary audit.",
        "- `figures/transfer_failure_attribution.*`: Figure 7, compact-site spatial guard audit for EIC transfer failures.",
        "- `figures/domain_support_audit.*`: Figure 8, train-side applicability and domain-support audit.",
        "- `figures/journal_readiness_audit.*`: Figure 9, CG/EG readiness and claim-boundary audit.",
        "- `figures/coordinate_label_coverage_audit.*`: Figure 10, public coordinate-label coverage audit for bounded EG-readiness.",
        "- `figures/voi_backtest_audit.*`: Figure 11, retrospective full-field VOI backtest for bounded observation-design readiness.",
        "- `source_data/*_source_data.csv`: plotted source data for the audited main figure set where source-data tables are generated.",
        "- `figures/arcticdata_wedge_operating_curve.*`: additional wedge-ice recall/precision operating-curve audit.",
        "- `figures/synthetic_rare_cryostructure_audit.*`: additional synthetic high-EIC and rare-facies operating-point audit.",
        "- `figures/rare_facies_hybrid_operating_curve.*`: additional synthetic rare-facies hybrid operating-point audit.",
        "- `figures/posterior_uncertainty_alignment.*`: additional posterior uncertainty-error alignment audit.",
        "- `figures/computational_footprint_summary.*`: additional parameter, artifact-size and posterior-sample footprint audit.",
        "- `supplementary_figures/*`: duplicate figure-atlas copies retained for package audit; the authoritative submission figure directory is `figures/`.",
        "",
        "## Evidence Gate Snapshot",
        "",
        f"- Independent public sources passed: `{gate.get('independent_public_sources_passed', 'not available')}`.",
        f"- Passed validation tasks: `{gate.get('passed_tasks', 'not available')}` / `{gate.get('total_tasks', 'not available')}`.",
        f"- EIC sources passed: `{gate.get('eic_sources_passed', 'not available')}`.",
        f"- Cryofacies sources passed: `{gate.get('facies_sources_passed', 'not available')}`.",
        f"- CG model evidence passed: `{gate.get('cg_model_evidence_passed', 'not available')}`.",
        "",
        "## Boundary Notes",
        "",
        "- Jago River 2018 is used as a third independent targeted EIC source, not as a regional benchmark.",
        "- Wedge-ice handling is recall-oriented; the operating-curve audit exposes false-positive trade-offs rather than hiding them.",
        "- Synthetic rare-cryostructure handling is separated into high-EIC event screening and explicit rare-facies recall, so mean IoU is not treated as sufficient evidence.",
        "- The rare-facies hybrid is an explicit synthetic operating point for wedge recall, not a replacement for the main diffusion posterior.",
        "- Posterior uncertainty is evaluated as an error-localization diagnostic, not only as interval coverage.",
        "- Computational footprint is reported alongside accuracy, so compact and high-parameter posterior variants are not treated as cost-equivalent.",
        "- Transfer-failure attribution is diagnostic: it shows how the compact-site spatial guard controls exposed EIC transfer failures under per-site baselines without converting the paper into a separate application study.",
        "- Public field releases validate partial labels and proxies, whereas full-field volumetric validation remains synthetic.",
        "- VOI-ranked boreholes and ERT lines now have a synthetic full-field retrospective backtest, but they are still not prospectively field-optimized acquisition plans.",
        "- The package supports a CG-style computational geoscience algorithm manuscript; a stronger EG-style regional field claim would require prospective validation and public full-field 3D ground truth.",
        "",
        "## Reproducibility Links",
        "",
        "- `tables/real_data_cg_gate.json`: public-data evidence gate summary.",
        "- `tables/figure_atlas.csv`: machine-readable atlas for the algorithm/public-data figures copied into this submission package.",
        "- `tables/innovation_positioning_audit.csv`: innovation dimensions mapped to evidence coverage, maturity and claim boundaries.",
        "- `tables/innovation_positioning_summary.json`: machine-readable innovation-positioning summary and boundary note.",
        "- `tables/voi_backtest_audit.csv`: retrospective synthetic full-field VOI error-enrichment and ranked-target audit.",
        "- `tables/voi_backtest_summary.json`: machine-readable VOI backtest readiness summary and boundary note.",
        "- `tables/coordinate_label_coverage_audit.csv`: site-wise public coordinate and label-density audit for EG-readiness.",
        "- `tables/coordinate_label_coverage_summary.json`: machine-readable coordinate-label coverage summary and boundary note.",
        "- `tables/domain_support_site_audit.csv`: train-side domain-support scores, applicability classes and holdout outcomes by public site.",
        "- `tables/domain_support_summary.json`: machine-readable domain-support summary and applicability boundary.",
        "- `tables/journal_readiness_audit.csv`: CG/EG readiness criteria separating algorithm evidence from field-generalization gaps.",
        "- `tables/journal_readiness_summary.json`: machine-readable readiness summary and recommended positioning.",
        "- `tables/real_data_cg_benchmark.csv`: task-level public validation metrics.",
        "- `tables/public_data_token_inventory.csv`: processed observation-token inventory.",
        "- `tables/public_data_provenance.csv`: source URL, local processing and hash provenance.",
        "- `tables/site_investigation_boreholes.csv`: ranked additional borehole targets.",
        "- `tables/site_investigation_ert_lines.csv`: ranked additional ERT survey lines.",
        "- `tables/arcticdata_wedge_operating_points.csv`: wedge-ice operating-point comparison.",
        "- `tables/arcticdata_wedge_operating_curve.csv`: threshold-wise wedge-ice recall/precision curve.",
        "- `tables/arcticdata_wedge_probability_holdout_scores.csv`: hold-out wedge probability scores used by the curve.",
        "- `tables/synthetic_rare_cryostructure_audit.csv`: high-EIC and rare cryostructure metrics under raw and observation-rate-constrained operating points.",
        "- `tables/diffusion_rare_facies_hybrid_metrics.csv`: synthetic physics-trained, implicit-proposal and rare-facies hybrid metrics.",
        "- `tables/rare_facies_hybrid_operating_curve.csv`: EIC-floor sweep for the synthetic rare-facies hybrid.",
        "- `tables/posterior_uncertainty_alignment.csv`: uncertainty-error rank alignment and top-uncertainty error enrichment by posterior model and target.",
        "- `tables/computational_footprint.csv`: model parameter count, checkpoint size, prediction footprint, posterior samples and training-history metadata.",
        "- `tables/external_generalization_site_deltas.csv`: site-wise ArcticData holdout deltas for facies, EIC and wedge recall.",
        "- `tables/external_generalization_audit.csv`: task-level external generalization summary with site win rates and failure sites.",
        "- `tables/transfer_failure_site_diagnostics.csv`: site-wise EIC transfer outcomes, adaptive EIC methods, guard reasons and attribution labels.",
        "- `tables/transfer_failure_attribution_summary.csv`: exploratory small-n transfer-attribution associations and reason counts.",
        "",
    ]
    out = package_dir / "README.md"
    out.write_text("\n".join(lines), encoding="utf-8")
    return out


def _write_zip(package_dir: Path, package_zip: Path) -> Path:
    if package_zip.exists():
        package_zip.unlink()
    with zipfile.ZipFile(package_zip, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in _iter_package_files(package_dir, package_zip):
            zf.write(path, path.relative_to(package_dir).as_posix())
    return package_zip


def make_submission_package(root: Path) -> SubmissionPackageResult:
    root = root.resolve()
    paper_dir = root / "paper"
    figure_dir = root / "outputs" / "figures"
    source_dir = root / "outputs" / "source_data"
    table_dir = root / "outputs" / "tables"
    article_md = paper_dir / "cold_recon_nature_article.md"
    article_docx = paper_dir / "cold_recon_nature_article.docx"

    if not article_md.exists():
        raise FileNotFoundError(f"Missing article Markdown: {article_md}")

    build_article_docx(article_md, article_docx)

    package_dir = paper_dir / "submission_package"
    if package_dir.exists():
        shutil.rmtree(package_dir)
    package_dir.mkdir(parents=True, exist_ok=True)

    _copy_file(article_md, package_dir / article_md.name)
    _copy_file(article_docx, package_dir / article_docx.name)
    for name in ("cold_recon_claim_evidence_audit.md", "cold_recon_manuscript_draft.md", "references.bib"):
        src = paper_dir / name
        if src.exists():
            _copy_file(src, package_dir / name)
    _copy_supplementary_figure_atlas(root, package_dir)

    for stem in FIGURE_STEMS:
        for ext in FIGURE_EXTENSIONS:
            src = figure_dir / f"{stem}{ext}"
            if src.exists():
                _copy_file(src, package_dir / "figures" / src.name)
    _copy_all_submission_figures(root, package_dir)

    for name in SOURCE_DATA_FILES:
        src = source_dir / name
        if src.exists():
            dst = package_dir / "source_data" / src.name
            if name == "nature_figure_4_source_data.csv":
                _copy_csv_with_algorithm_labels(src, dst)
            else:
                _copy_file(src, dst)

    for name in TABLE_FILES:
        src = table_dir / name
        if src.exists():
            dst = package_dir / "tables" / src.name
            if name == "figure_atlas.csv":
                _copy_submission_figure_atlas_csv(src, dst)
            elif name == "site_investigation_boreholes.csv":
                _copy_csv_with_algorithm_labels(src, dst)
            else:
                _copy_file(src, dst)

    package_readme = _write_package_readme(root, package_dir)
    package_zip = _write_zip(package_dir, package_dir / "cold_recon_submission_package.zip")
    return SubmissionPackageResult(
        article_docx=article_docx,
        package_dir=package_dir,
        package_readme=package_readme,
        package_zip=package_zip,
    )
